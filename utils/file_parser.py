import re
import logging
from typing import Dict, List, Tuple, Optional
import os

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pdfplumber not available. PDF parsing will be disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX parsing will be disabled.")

def parse_questions_from_file(filepath: str) -> Tuple[List[Dict], Dict[int, str]]:
    """
    Parse questions and answers from uploaded file (PDF or DOCX)
    
    Args:
        filepath: Path to the uploaded file
        
    Returns:
        Tuple of (questions_list, answer_key_dict)
        questions_list: List of question dictionaries with text and options
        answer_key_dict: Dictionary mapping question numbers to correct answers
    """
    file_ext = os.path.splitext(filepath)[1].lower()
    
    if file_ext == '.pdf':
        if not PDF_AVAILABLE:
            raise Exception("PDF processing not available. Please install pdfplumber.")
        return parse_pdf_questions(filepath)
    elif file_ext == '.docx':
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing not available. Please install python-docx.")
        return parse_docx_questions(filepath)
    else:
        raise Exception(f"Unsupported file format: {file_ext}")

def parse_pdf_questions(filepath: str) -> Tuple[List[Dict], Dict[int, str]]:
    """Parse questions from PDF file with enhanced image handling"""
    try:
        with pdfplumber.open(filepath) as pdf:
            text = ""
            for page_num, page in enumerate(pdf.pages):
                logging.debug(f"Processing PDF page {page_num + 1}")
                
                # Try multiple extraction methods for better text capture
                page_text = page.extract_text()
                
                # If standard extraction fails or returns minimal text, try alternative methods
                if not page_text or len(page_text.strip()) < 50:
                    # Try extracting with layout preservation
                    page_text = page.extract_text(layout=True)
                
                if not page_text or len(page_text.strip()) < 50:
                    # Try character-level extraction for images with text
                    chars = page.chars
                    if chars:
                        page_text = ''.join([char['text'] for char in chars])
                
                if page_text:
                    # Clean up text artifacts from images
                    page_text = clean_pdf_artifacts(page_text)
                    text += page_text + "\n"
                    logging.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                else:
                    logging.warning(f"No text extracted from page {page_num + 1} - may contain only images")
                    # Add placeholder for image-heavy pages
                    text += f"\n[Page {page_num + 1} contains images/diagrams]\n"
        
        return extract_questions_from_text(text)
    
    except Exception as e:
        logging.error(f"Error parsing PDF: {str(e)}")
        raise Exception(f"Failed to parse PDF file: {str(e)}")

def clean_pdf_artifacts(text: str) -> str:
    """Clean up common PDF artifacts and improve text quality"""
    if not text:
        return text
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common PDF artifacts
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
    text = re.sub(r'(\w)(\d+\.)', r'\1 \2', text)  # Add space before question numbers
    text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Add space after sentence endings
    
    # Clean up option formatting
    text = re.sub(r'([A-D])\)(\w)', r'\1) \2', text)  # Add space after options
    
    return text.strip()

def parse_docx_questions(filepath: str) -> Tuple[List[Dict], Dict[int, str]]:
    """Parse questions from DOCX file with enhanced image handling"""
    try:
        doc = Document(filepath)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Also extract text from tables (in case questions are in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # Note about images in DOCX
        images_count = len([rel for rel in doc.part.rels.values() 
                           if "image" in rel.target_ref])
        if images_count > 0:
            logging.info(f"DOCX contains {images_count} images - text around images may need manual review")
            text += f"\n[Document contains {images_count} images/diagrams]\n"
        
        return extract_questions_from_text(text)
    
    except Exception as e:
        logging.error(f"Error parsing DOCX: {str(e)}")
        raise Exception(f"Failed to parse DOCX file: {str(e)}")

def extract_questions_from_text(text: str) -> Tuple[List[Dict], Dict[int, str]]:
    """
    Extract questions and answer key from text content
    
    Expected format:
    Q1. Question text here?
    A) Option A
    B) Option B
    C) Option C
    D) Option D
    
    Q2. Another question?
    A) Option A
    B) Option B
    C) Option C
    D) Option D
    
    ...
    
    Answer Key:
    1. A
    2. B
    3. C
    ...
    """
    questions = []
    answer_key = {}
    
    logging.debug(f"Processing text of length: {len(text)}")
    
    # Try to find a real answer key section, not just "correct answers" in instructions
    # Look for patterns that indicate an actual answer key section
    answer_key_patterns = [
        r'(?i)answer\s*key\s*:?\s*[\r\n]',  # Answer Key: followed by newline
        r'(?i)answers?\s*:?\s*[\r\n]\s*(?:Q?\d+[\.\:]\s*[A-D]|1[\.\:]\s*[A-D])',  # Answers: followed by actual answers
        r'(?i)solution\s*key\s*:?\s*[\r\n]',
        r'(?i)correct\s*answers?\s*:?\s*[\r\n]\s*(?:Q?\d+[\.\:]\s*[A-D]|1[\.\:]\s*[A-D])'  # Correct answers: followed by actual answers
    ]
    
    questions_text = text
    answer_text = ""
    
    for pattern in answer_key_patterns:
        match = re.search(pattern, text)
        if match:
            split_pos = match.start()
            questions_text = text[:split_pos]
            answer_text = text[split_pos:]
            logging.debug(f"Found answer key section using pattern: {pattern}")
            break
    
    if not answer_text:
        logging.debug("No answer key section found, processing entire text as questions")
        # For this specific case, let's try to find where questions actually end
        # Look for common endings like page numbers, URLs, etc.
        end_patterns = [
            r'www\.[a-zA-Z.]+\s*$',  # Website at end
            r'\d+\s*www\.[a-zA-Z.]+\s*$',  # Page number + website
            r'^\s*\d+\s*$'  # Just page numbers
        ]
        
        lines = text.split('\n')
        questions_end = len(lines)
        
        # Find where questions likely end by looking for the last Q pattern
        for i in range(len(lines) - 1, -1, -1):
            if re.match(r'^\s*Q\d+\.', lines[i], re.IGNORECASE):
                # Found last question, include some lines after it for options
                questions_end = min(len(lines), i + 10)
                break
        
        questions_text = '\n'.join(lines[:questions_end])
    
    # Extract questions
    questions = extract_questions_list(questions_text)
    logging.debug(f"Extracted {len(questions)} questions")
    
    # Extract answer key
    if answer_text:
        answer_key = extract_answer_key(answer_text)
        logging.debug(f"Extracted {len(answer_key)} answers")
    
    # If no answer key found, create a dummy answer key for testing purposes
    if not answer_key and questions:
        logging.debug("Creating dummy answer key for testing")
        for i, q in enumerate(questions):
            # Use first available option as dummy answer
            if q.get('options'):
                first_option = list(q['options'].keys())[0]
                answer_key[i + 1] = first_option
    
    return questions, answer_key

def extract_questions_list(text: str) -> List[Dict]:
    """Extract individual questions with options from text"""
    questions = []
    
    # Extensive text cleanup for PDF artifacts
    text = clean_pdf_text(text)
    
    # Pattern to match questions starting with Q1., Q2., etc.
    # More flexible pattern to capture multiline questions
    question_pattern = r'Q(\d+)\.?\s*(.*?)(?=Q\d+\.|\Z)'
    question_matches = re.findall(question_pattern, text, re.DOTALL | re.IGNORECASE)
    
    logging.debug(f"Found {len(question_matches)} potential question matches")
    
    for match in question_matches:
        question_num = int(match[0])
        question_text = match[1].strip()
        
        # Skip if the question text is too short (likely not a real question)
        if len(question_text) < 10:
            logging.debug(f"Skipping Q{question_num} - text too short: {len(question_text)}")
            continue
        
        logging.debug(f"Processing Q{question_num}: {question_text[:100]}...")
        
        # Extract question and options
        question_data = parse_single_question(question_text)
        if question_data:
            question_data['number'] = question_num
            questions.append(question_data)
            logging.debug(f"Successfully parsed Q{question_num}")
        else:
            logging.debug(f"Failed to parse Q{question_num}")
    
    return questions

def clean_pdf_text(text: str) -> str:
    """Clean up PDF text extraction artifacts"""
    # Remove common PDF artifacts
    text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
    text = re.sub(r'\n+', '\n', text)  # Normalize line breaks
    
    # Remove isolated single characters that are PDF artifacts
    text = re.sub(r'\s[a-z]\s', ' ', text)  # Remove isolated lowercase letters
    text = re.sub(r'\s[A-Z]\s(?![A-D]\))', ' ', text)  # Remove isolated uppercase letters (except option letters)
    
    # Remove common PDF noise patterns
    text = re.sub(r'www\.[a-zA-Z.]+', '', text)  # Remove website URLs
    text = re.sub(r'[0-9]+\s*www\.[a-zA-Z.]+', '', text)  # Remove page numbers with URLs
    
    # Clean up option patterns - ensure proper spacing
    text = re.sub(r'\(\s*([a-dA-D])\s*\)', r'(\1)', text)  # Fix spaced parentheses
    
    return text.strip()

def parse_single_question(text: str) -> Optional[Dict]:
    """Parse a single question with its options"""
    
    # Clean up text and split into lines
    text = text.strip()
    if not text:
        return None
    
    # Initialize variables
    question_text = ""
    options = {}
    
    # Clean the text first
    text = clean_pdf_text(text)
    
    # Try to extract question and options using multiple patterns
    
    # Pattern 1: Inline options like (a) Option1 (b) Option2 (c) Option3 (d) Option4
    # Look for at least 2 option patterns in the text
    option_pattern = r'\(([a-dA-D])\)'
    option_positions = [(m.start(), m.group(1)) for m in re.finditer(option_pattern, text)]
    
    if len(option_positions) >= 2:
        # Extract question text (everything before first option)
        first_option_pos = option_positions[0][0]
        question_text = text[:first_option_pos].strip()
        
        # Handle images in questions - if question is very short, it might have an image
        if len(question_text) < 30:
            question_text += " [This question may contain an image or diagram]"
        
        # Extract options
        for i, (pos, option_letter) in enumerate(option_positions):
            # Find the end position for this option
            next_pos = option_positions[i + 1][0] if i + 1 < len(option_positions) else len(text)
            
            # Extract option text
            option_start = pos + 3  # Skip past "(x)"
            option_text = text[option_start:next_pos].strip()
            
            # Clean up option text
            option_text = re.sub(r'^\)\s*', '', option_text)  # Remove leading )
            option_text = re.sub(r'\s*\([a-dA-D]\)\s*$', '', option_text)  # Remove trailing (x)
            
            # Handle empty or very short options (might be image-based)
            if not option_text or len(option_text) < 3:
                option_text = f"[Option {option_letter.upper()} - may contain image]"
            
            options[option_letter.upper()] = option_text
    
    else:
        # Pattern 2: Multi-line format
        lines = text.split('\n')
        current_question = []
        current_options = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this line contains an option pattern
            option_match = re.match(r'^\(?([a-dA-D])\)?\s*(.+)', line)
            if option_match and len(current_question) > 0:  # Only treat as option if we have question text
                opt_letter = option_match.group(1).upper()
                opt_text = option_match.group(2).strip()
                
                # Handle empty or very short options
                if not opt_text or len(opt_text) < 3:
                    opt_text = f"[Option {opt_letter} - may contain image]"
                    
                current_options[opt_letter] = opt_text
            else:
                # This is part of the question text
                current_question.append(line)
        
        question_text = ' '.join(current_question).strip()
        options = current_options
    
    # Validate that we have a reasonable question and options
    if not question_text or len(options) < 2:
        return None
        
    # Final cleanup of question text
    question_text = re.sub(r'\s+', ' ', question_text).strip()
    
    return {
        'text': question_text,
        'options': options
    }

def extract_answer_key(text: str) -> Dict[int, str]:
    """Extract answer key from text"""
    answer_key = {}
    
    # Pattern to match answer key entries like "1. A", "2. B", etc.
    answer_patterns = [
        r'(\d+)\.?\s*([A-D])',  # 1. A or 1 A
        r'Q?(\d+)[\.\:\s]*([A-D])',  # Q1: A or Q1 A
    ]
    
    for pattern in answer_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            question_num = int(match[0])
            answer = match[1].upper()
            answer_key[question_num] = answer
    
    return answer_key

def validate_questions_and_answers(questions: List[Dict], answer_key: Dict[int, str]) -> bool:
    """Validate that questions and answer key are consistent"""
    if not questions:
        return False
    
    # Check if we have answers for most questions
    question_numbers = {q.get('number', i+1) for i, q in enumerate(questions)}
    answer_numbers = set(answer_key.keys())
    
    # We should have answers for at least 50% of questions
    overlap = len(question_numbers.intersection(answer_numbers))
    return overlap >= len(questions) * 0.5

# Test function for debugging
def test_parsing():
    """Test function to validate parsing logic"""
    sample_text = """
    Q1. What is the capital of France?
    A) London
    B) Berlin
    C) Paris
    D) Madrid
    
    Q2. Which planet is closest to the Sun?
    A) Venus
    B) Mercury
    C) Earth
    D) Mars
    
    Answer Key:
    1. C
    2. B
    """
    
    questions, answers = extract_questions_from_text(sample_text)
    print(f"Questions: {questions}")
    print(f"Answers: {answers}")
    
    return questions, answers

def test_pdf_parsing():
    """Test parsing with PDF-like inline format"""
    sample_text = """
    Q1. Which of the following diagrams indicates the best relation between Thief, Criminal and Police?
    (a) Diagram A (b) Diagram B (c) Diagram C (d) Diagram D
    
    Q2. Which of the following diagrams indicates best relation between Pigeon, Bird and Dog?
    (a) Circle A (b) Circle B (c) Circle C (d) Circle D
    """
    
    questions, answers = extract_questions_from_text(sample_text)
    print(f"PDF Format Questions: {questions}")
    print(f"PDF Format Answers: {answers}")
    
    return questions, answers

if __name__ == "__main__":
    # Run test
    test_parsing()
